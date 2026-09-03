"""Deliverable generation. Turns agent output into files a human would sign.

python-docx is used rather than any AI library. This module contains no
model calls at all; it takes structured content and formats it.

Always render and LOOK at the output before shipping a template change:
    soffice --headless --convert-to pdf note.docx && pdftoppm -jpeg -r 100 note.pdf p
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core import trace

OUT = Path(__file__).resolve().parents[2] / "outputs"


def approval_note(title, reference, background, findings, recommendation,
                  prepared_by="", filename="approval_note.docx"):
    """findings: list of dicts with keys item, observation, severity."""
    OUT.mkdir(exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Reference: {reference}").italic = True

    doc.add_heading("1. Background", level=2)
    doc.add_paragraph(background)

    doc.add_heading("2. Findings", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Item", "Observation", "Severity"]):
        cell.paragraphs[0].add_run(text).bold = True
    for f in findings:
        row = table.add_row().cells
        row[0].text = str(f.get("item", ""))
        row[1].text = str(f.get("observation", ""))
        row[2].text = str(f.get("severity", ""))
    for row in table.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(4.0)
        row.cells[2].width = Inches(1.1)

    doc.add_heading("3. Recommendation", level=2)
    doc.add_paragraph(recommendation)

    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Prepared by: ").bold = True
    sig.add_run(prepared_by or "________________")
    doc.add_paragraph("Approved by: ________________        Date: ____________")

    path = OUT / filename
    doc.save(path)
    trace.emit("deliverable", kind="docx", path=str(path), findings=len(findings))
    return str(path)
